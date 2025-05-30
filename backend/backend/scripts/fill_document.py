from docx import Document
from docxtpl import DocxTemplate
import os
from workalendar.europe import Russia
import locale
from datetime import date
from random import randint
import re
from django.core.exceptions import ValidationError

from core.settings.base import DOCUMENTS_FOLDER, TEMPLATES_FOLDER, MEDIA_ROOT


def fill_document(filename: str, data: dict, table_data: list[list[str]]) -> dict:
    """
    Выполняет заполнение документа шаблона `filename` необходимыми данными. 
    Итоговый файл сохранятся в `DOCUMENTS_FOLDER`.
    Заполнение шаблонов происходит согласно подстановке значений по необходимым ключам в плейсхолдеры в шаблоне.
    Заполнение таблицы происходит с помощью добавления строк в таблицу согласно данным в списке списков `table_data`.
    
    ВАЖНО! 
    - Все строки добавляются "как есть", поэтому важно соблюдать порядок элементов в списке списков.
    - Шаблон должен быть настроен на добавление таблиц, имея единственную строку с заполненными ячейками, начинающимися на "RC".
    
    :return: 
    Возвращает словарь результата: `{'code': int, 'error': str | None, 'path': str, 'warnings': list[str]}`.

    Code - внутренний код результата. 10 в случае успеха (файл сохранён как `templatename_%m-%y`), иначе:
    - 21 - файл не найден в папке шаблонов;
    - 22 - поле `order_date` не найдено в шаблоне документа;
    - 23 - не найдена строка для заполнения таблицы;
    - 24 - ...

    :param filename: Имя файла шаблона. Файл должен располагаться в `TEMPLATES_FOLDER`.
    :param data: Словарь, содержащий данные для заполнения. Предполагается, что ключи содержатся в самом документе.
    :param table_data: Список списков, содержащий данные для заполнения таблицы.

    :see: core.settings
    """

    def find_placeholders(template: DocxTemplate) -> set:
        """Находит все плейсхолдеры в документе"""
        placeholders = set()
        pattern = re.compile(r'\{\{.*?\}\}')

        for paragraph in template.get_docx().paragraphs:
            matches = pattern.findall(paragraph.text)
            placeholders.update(matches)
        
        for table in template.get_docx().tables:
            for row in table.rows:
                for cell in row.cells:
                    matches = pattern.findall(cell.text)
                    placeholders.update(matches)

            return sorted(placeholders)

    def find_empty_row(table):
        """Находит первую полностью пустую строку в таблице"""
        for i, row in enumerate(table.rows):
            if all(cell.text.startswith('RC') for cell in row.cells):
                return i
        return None  # Если пустой строки нет

    def copy_row_formatting(source_row, target_row):
        """Копирует все параметры форматирования из исходной строки в целевую"""
        for source_cell, target_cell in zip(source_row.cells, target_row.cells):
            # Копируем параметры абзаца
            source_paragraph = source_cell.paragraphs[0]
            target_paragraph = target_cell.paragraphs[0]
            
            # Шрифт
            source_font = source_paragraph.runs[0].font if source_paragraph.runs else None

            target_run = target_paragraph.add_run()
            if source_font:
                target_run.font.name = source_font.name
                target_run.font.size = source_font.size
                target_run.font.bold = source_font.bold
                target_run.font.italic = source_font.italic
                target_run.font.underline = source_font.underline
                target_run.font.color.rgb = source_font.color.rgb
            
            # Выравнивание
            target_paragraph.alignment = source_paragraph.alignment
            
            # Границы ячеек (если нужно)
            #target_cell._tc.tcPr = deepcopy(source_cell._tc.tcPr)

    result = {
        'code': 0,
        'error': None,
        'path': '',
        'warnings': [],
        'shown_date': None,
    }

    # Для русскоязычных дат
    locale.setlocale(locale.LC_ALL, 'ru_RU')

    # Сделать место сохранение документа
    if "/" in filename:
        filename = filename.split("/")[-1]

    # Получить первый рабочий день текущего месяца и текущего года
    data["order_date"] = Russia().add_working_days(date(date.today().year, date.today().month, 1), 0).strftime("%d %B %Y")
    result['shown_date'] = data["order_date"]

    if os.path.exists(TEMPLATES_FOLDER / filename) is False:
        result['code'] = 21
        result['error'] = f"Файл {TEMPLATES_FOLDER / filename} не найден."
        return result

    doc = DocxTemplate(TEMPLATES_FOLDER / filename)
    placeholders = find_placeholders(doc)
    if "{{ order_date }}" not in placeholders:
        result['code'] = 22
        result['error'] = f"В шаблоне {TEMPLATES_FOLDER / filename} должно быть поле `order_date` для указания даты, но оно не было найдено."
        return result
    for placeholder in placeholders:
        if placeholder not in data.keys():
            result['warnings'].append(f"В шаблоне {TEMPLATES_FOLDER / filename} найдено поле `{placeholder}`, но оно не было представлено в данных. Данный плейсхолдер не был заменён.")
    for placeholder in data.keys():
        if placeholder not in placeholders:
            result['warnings'].append(f"В шаблоне {TEMPLATES_FOLDER / filename} не найдено поле `{placeholder}`, но оно было представлено в данных. Данный плейсхолдер не был использован.")

    doc.render(data)

    # Сохранить в папку документах
    randcode = randint(0, 1000000)
    doc.save(f"tmp/~$tmp{randcode}.docx")

    # Работаем с таблицами

    doc = Document(f"tmp/~$tmp{randcode}.docx")
    table = doc.tables[0]
    empty_row_idx = find_empty_row(table)
    if empty_row_idx is None:
        os.remove(f"tmp/~$tmp{randcode}.docx")
        result['code'] = 23
        result['error'] = "В документе должна быть строка для заполнения таблицы, но она не была найдена."

    empty_row = table.rows[empty_row_idx]

    table._tbl.remove(empty_row._element)

    for row_data in table_data:
        new_row = table.add_row()
        table.rows[empty_row_idx]._element.addprevious(new_row._element)
        
        copy_row_formatting(empty_row, new_row)
        
        # Заполняем данными
        for col_idx, value in enumerate(row_data):
            new_row.cells[col_idx].paragraphs[0].runs[0].text = str(value)
        
        empty_row_idx += 1

    os.remove(f"tmp/~$tmp{randcode}.docx")
    # TODO: сохранять в директории комании - получать название компании в data
    result['code'] = 10
    result['path'] = DOCUMENTS_FOLDER / f"{filename}_{date(date.today().year, date.today().month, 1).strftime('%m-%y')}.docx"
    doc.save(result['path'])
    
    return result


def find_fields(filename) -> list:
    """Находит все плейсхолдеры в документе"""

    def reformat_placeholder_names(placeholders: set) -> list:
        """Форматирует имена плейсхолдеров для использования в шаблоне"""
        if placeholders is None:
            return []
        return [placeholder.replace("{{", "").replace("}}", "").strip() for placeholder in placeholders]
    
    def find_placeholders(template: DocxTemplate) -> set:
        """Находит все плейсхолдеры в документе"""
        # if not template or not hasattr(template, 'docx') or template.docx is None:
        #     print(template.get_docx())
        #     raise ValidationError("Не удалось загрузить документ или документ поврежден")

        placeholders = set()
        pattern = re.compile(r'\{\{.*?\}\}')

        for paragraph in template.get_docx().paragraphs:
            matches = pattern.findall(paragraph.text)
            placeholders.update(matches)
        
        for table in template.get_docx().tables:
            for row in table.rows:
                for cell in row.cells:
                    matches = pattern.findall(cell.text)
                    placeholders.update(matches)

        return sorted(placeholders)
    
    doc = DocxTemplate(MEDIA_ROOT / filename.name) #! Теперь файлы загружать именно так!
    placeholders = find_placeholders(doc)
    reformated = reformat_placeholder_names(placeholders)

    PREORDER_FIELDS = [
        "contract_number",
        "contract_date",
        "order_date",
        "order_number",
        "contrator_person",
        "contractor_company",
        "contractor_post",
        "contractor_company_full",
        "executor_person",
        "executor_post",
        "executor_company_full",
        "executor_company",
    ]
    return [placeholder for placeholder in reformated if placeholder not in PREORDER_FIELDS]
